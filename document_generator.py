import os
import uuid
import qrcode
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime 
from document_register import GoogleDriveManager

class DocumentGenerator:
    def __init__(self, parent=None, drive_manager=None):
        """
        Inicializa la ventana del generador de documentos.
        
        Args:
            parent: Ventana padre (ventana del menú principal).
                Si no se proporciona, el generador funciona como aplicación independiente.
            drive_manager: Instancia de GoogleDriveManager ya inicializada
        """
        # Guardar referencia a la ventana padre y drive_manager
        self.parent = parent
        self.drive_manager = drive_manager
        
        # Crear la ventana principal
        self.window = ttk.Toplevel(parent) if parent else ttk.Window()
        self.window.title("Generación de Documentos")
        
        # Configurar tamaño y posición
        window_width = 800
        window_height = 600
        
        # Centrar en pantalla
        screen_width = self.window.winfo_screenwidth()
        screen_height = self.window.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        # Establecer geometría
        self.window.geometry(f'{window_width}x{window_height}+{x}+{y}')
        
        if parent:
            self.window.transient(parent)
            self.window.grab_set()
            
        # Configurar el resto de la interfaz
        self.setup_ui()
        
    def setup_ui(self):
        """Configura la interfaz de usuario del generador"""
        # Frame principal con padding
        main_frame = ttk.Frame(self.window, padding=20)
        main_frame.pack(fill=BOTH, expand=YES)
        
        # Título
        title_label = ttk.Label(
            main_frame,
            text="Generación de Documentos",
            font=("Segoe UI", 24),
            bootstyle="primary"
        )
        title_label.pack(pady=(0, 30))
        
        # Frame para el formulario
        form_frame = ttk.LabelFrame(
            main_frame,
            text="Configuración",
            padding=20
        )
        form_frame.pack(fill=X, pady=20)
        
        # Frame para la entrada de cantidad
        cantidad_frame = ttk.Frame(form_frame)
        cantidad_frame.pack(fill=X, pady=10)
        
        # Label para la cantidad
        cantidad_label = ttk.Label(
            cantidad_frame,
            text="Número de actas a generar:",
            font=("Segoe UI", 12)
        )
        cantidad_label.pack(side=LEFT)
        
        # Entry para la cantidad
        self.cantidad_var = tk.StringVar()
        cantidad_entry = ttk.Entry(
            cantidad_frame,
            textvariable=self.cantidad_var,
            font=("Segoe UI", 12),
            width=10
        )
        cantidad_entry.pack(side=LEFT, padx=10)

        # Vincular Enter al botón generar
        cantidad_entry.bind('<Return>', lambda event: self.iniciar_generacion())

        
        # Frame para los botones
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=30)
        
        # Botón generar
        generar_btn = ttk.Button(
            button_frame,
            text="Generar Actas",
            command=self.iniciar_generacion,
            bootstyle="primary",
            width=20
        )
        generar_btn.pack(side=LEFT, padx=5)
        
        # Botón volver
        volver_btn = ttk.Button(
            button_frame,
            text="Volver al Menú",
            command=self.volver_menu,
            bootstyle="danger-outline",
            width=20
        )
        volver_btn.pack(side=LEFT, padx=5)
            
    def generar_qr(self, uuid_texto, ruta_salida):
        """
        Genera un código QR y lo guarda como imagen.
        
        Args:
            uuid_texto (str): El texto (UUID) que se codificará en el QR
            ruta_salida (str): Ruta donde se guardará la imagen del QR
        """
        try:
            # Configurar el código QR
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            
            # Agregar los datos y generar el QR
            qr.add_data(uuid_texto)
            qr.make(fit=True)
            
            # Crear y guardar la imagen
            img = qr.make_image(fill_color="black", back_color="white")
            img.save(ruta_salida)
            
            print(f"QR generado y guardado en: {ruta_salida}")
            return True
            
        except Exception as e:
            print(f"Error al generar QR: {str(e)}")
            return False
        
    def insertar_qr_en_tabla(self, doc, ruta_qr, numero_tabla, fila, columna, ancho=1, alto=1):
        """
        Inserta un QR en una celda de una tabla específica.
        
        Args:
            doc: Documento Word donde se insertará el QR
            ruta_qr (str): Ruta del archivo de imagen QR
            numero_tabla (int): Índice de la tabla en el documento (0-based)
            fila (int): Número de fila en la tabla
            columna (int): Número de columna en la tabla
            ancho (int): Ancho de la imagen en pulgadas
            alto (int): Alto de la imagen en pulgadas
        
        Returns:
            bool: True si la operación fue exitosa, False en caso contrario
        """
        try:
            # Obtener la tabla especificada
            tabla = doc.tables[numero_tabla]
            celda = tabla.rows[fila].cells[columna]

            # Eliminar contenido previo en la celda
            for parrafo in celda.paragraphs:
                celda._element.remove(parrafo._element)

            # Insertar la imagen en la celda
            parrafo = celda.add_paragraph()
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = parrafo.add_run()
            run.add_picture(ruta_qr, width=Inches(ancho), height=Inches(alto))

            print(f"QR insertado en tabla {numero_tabla}, celda ({fila}, {columna}).")
            return True
            
        except IndexError:
            print(f"Error: La tabla o la celda especificada ({fila}, {columna}) no existe.")
            return False
        except Exception as e:
            print(f"Error al insertar QR en tabla: {str(e)}")
            return False
        
    def generar_actas(self, cantidad, ruta_salida, numero_tabla, fila, columna):
        """
        Genera múltiples actas con QR únicos.
        
        Args:
            cantidad (int): Número de actas a generar
            ruta_salida (str): Carpeta donde se guardarán las actas
            numero_tabla (int): Índice de la tabla donde se insertará el QR
            fila (int): Fila de la tabla donde se insertará el QR
            columna (int): Columna de la tabla donde se insertará el QR
            
        Returns:
            bool: True si todas las actas se generaron correctamente, False en caso contrario
        """
        try:
            plantilla = "plantilla.docx"
            if not os.path.exists(plantilla):
                messagebox.showerror("Error", f"No se encontró la plantilla '{plantilla}'.")
                return False

            progress_window = ttk.Toplevel(self.window)
            progress_window.title("Generando Actas")
            progress_window.geometry("300x150")
            progress_window.transient(self.window)
            x = self.window.winfo_x() + (self.window.winfo_width() - 300) // 2
            y = self.window.winfo_y() + (self.window.winfo_height() - 150) // 2
            progress_window.geometry(f"+{x}+{y}")
            
            progress_frame = ttk.Frame(progress_window, padding=20)
            progress_frame.pack(fill=BOTH, expand=YES)
            
            progress_label = ttk.Label(
                progress_frame,
                text="Generando actas...",
                font=("Segoe UI", 10)
            )
            progress_label.pack(pady=(0, 10))
            
            progress_bar = ttk.Progressbar(
                progress_frame,
                length=200,
                mode='determinate'
            )
            progress_bar.pack()

            for i in range(1, cantidad + 1):
                progress = (i / cantidad) * 100
                progress_bar['value'] = progress
                progress_label.config(text=f"Generando acta {i} de {cantidad}")
                progress_window.update()

                uuid_texto = str(uuid.uuid4())
                ruta_qr = f"qr_acta_{i}.png"
                
                if not self.generar_qr(uuid_texto, ruta_qr):
                    raise Exception(f"Error al generar QR para acta {i}")

                # Registrar QR en el spreadsheet
                try:
                    self.drive_manager.add_qr_record(
                        qr_id=uuid_texto,
                        formato="GCO-REG-099",
                        fecha=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    )
                except Exception as e:
                    print(f"Error al registrar QR en spreadsheet: {str(e)}")

                doc = Document(plantilla)
                if not self.insertar_qr_en_tabla(doc, ruta_qr, numero_tabla, fila, columna):
                    raise Exception(f"Error al insertar QR en acta {i}")

                salida_docx = os.path.join(ruta_salida, f"acta_{i}.docx")
                doc.save(salida_docx)
                print(f"Acta guardada en: {salida_docx}")
                os.remove(ruta_qr)

            progress_window.destroy()
            messagebox.showinfo(
                "Éxito", 
                f"Se generaron {cantidad} actas correctamente en {ruta_salida}"
            )
            return True

        except Exception as e:
            print(f"Error al generar actas: {str(e)}")
            messagebox.showerror("Error", f"Ocurrió un error al generar las actas: {str(e)}")
            return False
        
    def iniciar_generacion(self):
        """
        Obtiene datos desde la GUI e inicia la generación de actas.
        Valida la entrada y gestiona el proceso completo de generación.
        """
        try:
            # Cambiar el cursor a "ocupado"
            self.window.config(cursor="wait")
            
            # Obtener y validar la cantidad
            try:
                cantidad = int(self.cantidad_var.get())
                if cantidad <= 0:
                    raise ValueError("La cantidad debe ser mayor a 0.")
            except ValueError as e:
                messagebox.showerror("Error", "Por favor, ingresa una cantidad válida mayor a 0.")
                self.window.config(cursor="")
                return

            # Solicitar carpeta de salida
            ruta_salida = filedialog.askdirectory(
                title="Selecciona la carpeta de salida",
                parent=self.window
            )
            
            if not ruta_salida:
                self.window.config(cursor="")
                return
                
            # Deshabilitar la interfaz mientras se generan las actas
            for widget in self.window.winfo_children():
                if isinstance(widget, (ttk.Button, ttk.Entry)):
                    widget.configure(state="disabled")
    
            try:
                # Valores fijos para la ubicación del QR en la tabla
                # Estos podrían ser configurables en una versión futura
                self.generar_actas(cantidad, ruta_salida, 1, 5, 0)
                
            except Exception as e:
                messagebox.showerror("Error", f"Ocurrió un error al generar las actas: {str(e)}")
                
            finally:
                # Restaurar la interfaz
                self.window.config(cursor="")
                for widget in self.window.winfo_children():
                    if isinstance(widget, (ttk.Button, ttk.Entry)):
                        widget.configure(state="normal")
                
        except Exception as e:
            messagebox.showerror("Error", f"Error inesperado: {str(e)}")
            self.window.config(cursor="")
        
    def volver_menu(self):
        """
        Cierra la ventana del generador y vuelve al menú principal.
        Si hay una generación en proceso, pide confirmación.
        """
        try:
            # Si hay una ventana padre (menú principal), mostrarla
            if self.parent:
                # Restaurar la ventana padre
                self.parent.deiconify()
                
                # Asegurarse de que la ventana padre esté activa
                self.parent.focus_force()
                
            # Cerrar la ventana actual
            self.window.destroy()
            
        except Exception as e:
            print(f"Error al volver al menú: {str(e)}")
            # En caso de error, intentar cerrar la ventana de todas formas
            self.window.destroy()

    def on_closing(self):
        """Maneja el cierre de la ventana"""
        if self.parent:
            self.parent.deiconify()  # Restaurar ventana principal
        self.root.destroy()
